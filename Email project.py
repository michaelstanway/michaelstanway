# -*- coding: utf-8 -*-
"""
Created on Mon Jun 28 13:32:37 2021

@author: micha
"""

#Email project

#Plan:
    #create data to read

    #code how to read that data and then code how to put that data into another file
    #code how to send that file to the desired email
    #code a way to run this on all systems(perhaps a website or application)

from docx import Document
import smtplib, ssl

myName = 'Michael' #Name of person sending the email
myEmail = "michaelsstanway@gmail.com" #Email of the person sending the email  
password = "urmumgay" #Password to the email being used to send
subject = 'Business' #Subject of the email
smtpServer = "smtp.gmail.com" #This will need to be changed depending on the email site used
port = 587
topic = "Subject: " + subject #Subject in the required useable format

client_name = [] #Empty list to store the names of the clients
client_company = [] #Empty list to store the names of the clients' respective companies
yourEmail = [] #Empty list to store the emails of the clients

def obtainClients(docFileName): #Function to read and save the clients
    document = Document(docFileName) #Opens the required document
    finalText = [] #Empty list to store data
    for line in document.paragraphs:
        finalText.append(line.text) #Saves the file to the list
    #return finalText
    return '/n'.join(finalText) #Joins up the list

def obtainText(docFileName): #Function to read and save the email template
    document = Document(docFileName) #Opens and reads the document
    finalText = []
    for line in document.paragraphs:
        finalText.append(line.text)
    #return finalText
    return ' '.join(finalText)

def message(docFileName, x, y, z): #Function to generate the message for each client
    lines = obtainText(docFileName) #Calls the function to access the template
    words = lines.split() #Breaks the template down into each word
    new_words = [] #Creates an empty list to store the individualised message
    for word in words:
        if word == 'clientname,': #Puts the client's name into the correct place
            word = x + ','
        elif word == 'company_1,': #Puts the company's name into the correct place
            word = y+','
        elif word == 'myname': #Puts the senders name into the correct place
            word = '\n' + z
        new_words.append(word) #Saves each word into the place for the individualised message
        
    final_message = new_words[0] #Cannot create an empty string as far as I know so the first word must be presaved
    for i in range(len(new_words)):
        if i>0: #So that the word already in the string is skipped
            word = new_words[i]
            if word == 'I': #As the first paragraph begins with I we must add a space here
                final_message = final_message + '\n' + '\n' + word
            elif word == 'Thanks': #Second paragraph
                final_message = final_message + '\n' + '\n' + word
            else:
                final_message = final_message + ' ' + word #Adds in all other words
    return(final_message)

def sendEmail(x, y): #Function to send emails
    context = ssl.create_default_context() #Not too sure what this means but its needed
    try: #Sends the message to desired place after accessing gmail server
        server = smtplib.SMTP(smtpServer, port)
        server.starttls(context = context)
        server.login(myEmail, password)
        for i in range(len(x)): #Ensures to send all required emails at once so that the code does not need to access the server each time
            server.sendmail(myEmail, y[i], x[i])
    except Exception as e: #In case the message cannot send for whatever reason
        print("the email could not be sent.")
    finally:
        server.quit() #Exits the server

clients = obtainClients('Clients.docx').split('/n') #Saves all the client information

for client in clients: #Looks at each client in the list
    if client != '': #In case any empty spaces have been saved
        client_info = client.split('\t') #Creates distinct names, companies, and email addresses
        client_name.append(client_info[0])#Creates a list of client names
        client_company.append(client_info[1])#Creates a list of company names
        yourEmail.append([client_info[2]])#Creates a list of contact email addresses

readyEmail = [] #Empty list to store the complete messages

for i in range(len(client_name)): #Generates and saves the correct messages for each client
    myEmail2 = "<" + myEmail + ">" #Email in a format that fits the use required
    recieverEmail = '<' + yourEmail[i][0] + '>' #Email in a format that fits the use required
    newEmail = "From: From " + myName + ' ' + myEmail2 + """
    To: To """ + client_name[i] + ' ' + recieverEmail +"""
    """ #Creates the first part of the message to be sent
    email = newEmail + '\n' + topic + '\n' + message('Emails.docx', client_name[i], client_company[i], myName) #Generates the full message by combining each part
    readyEmail.append(email) #Saves the message for each client

sendEmail(readyEmail, yourEmail) #Sends the final message